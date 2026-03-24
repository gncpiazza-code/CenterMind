from docx import Document
import os

doc = Document(r'C:\Users\cigar\OneDrive\Desktop\Agente Real Tabacalera\hoja_de_ruta_migracion.docx')

output_path = r'C:\Users\cigar\OneDrive\Desktop\BOT-SQL\antigravity\CenterMind\CenterMind\hoja_de_ruta_texto.txt'

with open(output_path, 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n\n")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            style = p.style.name if p.style else "None"
            f.write(f"[P{i}] [{style}]\n{text}\n\n")
    
    f.write("\n=== TABLES ===\n\n")
    for ti, t in enumerate(doc.tables):
        f.write(f"\n--- Table {ti}: {len(t.rows)} rows x {len(t.columns)} columns ---\n")
        for ri, r in enumerate(t.rows):
            cells = [c.text.strip() for c in r.cells]
            f.write(f"  Row {ri}:\n")
            for ci, cell in enumerate(cells):
                f.write(f"    Col {ci}: {cell}\n")
            f.write("\n")

print(f"Written to {output_path}")
print(f"File size: {os.path.getsize(output_path)} bytes")
