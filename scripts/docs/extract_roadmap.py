import docx
import os

doc_path = r'C:\Users\cigar\OneDrive\Desktop\Agente Real Tabacalera\hoja_de_ruta_migracion.docx'
output_path = r'C:\Users\cigar\OneDrive\Desktop\BOT-SQL\antigravity\CenterMind\roadmap_full.txt'

if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    text = '\n'.join([p.text for p in doc.paragraphs])
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"Success: Wrote {len(doc.paragraphs)} paragraphs to {output_path}")
else:
    print(f"Error: File not found at {doc_path}")
